from flask import Flask, request, render_template
import string
import requests
import bs4
from bs4.element import Comment
import csv
import config
from docx import Document


from sumy.parsers.html import HtmlParser
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer as Summarizer
from sumy.nlp.stemmers import Stemmer
from sumy.utils import get_stop_words
import nltk
nltk.download('punkt')



def MakeDoc(headers, text, useStuff, useStuffHead, uni, country):
    document = Document()
    for index, title in enumerate(headers):
        if text[index]:
            document.add_heading(title.capitalize(), 0)
            for item in text[index]:
                document.add_paragraph(str(item))
    
    for index, title in enumerate(useStuffHead):
        if useStuff[index]:
            document.add_heading(title.capitalize(), 0)
            if useStuffHead[index] == 'Sources':
                for item in useStuff[index]:
                    document.add_paragraph(str(item))
            elif useStuffHead[index] == 'Near You':
                for place in useStuff[index]:
                    document.add_paragraph(place)
                    for item in useStuff[index][place]:
                        if str(item) == 'nothing found':
                            document.add_paragraph(str(item))      
                        else:
                            for key in item:
                                document.add_paragraph(f'{key}: {item[key]}')
            elif useStuffHead[index] == 'Subjects' or useStuffHead[index] == 'Majors':
                if country == 'UK':
                    document.add_paragraph(f'Here is a link with a list of possible courses: {useStuff[index]}')
                elif country == 'US':
                    document.add_paragraph(f'Here is a link with a list of possible majors: {useStuff[index]}')
            else:
                document.add_paragraph(useStuff[index].replace((uni + ' University /'), ''))

    document.save('static\\' + uni + 'Research.docx')
    return document


def PyYelp(location):
    url = 'https://api.yelp.com/v3/businesses/search'
    headers = {
        "Authorization": "Bearer " + config.api_key
    }

    final = {
        'coffee shop': None,
        'Pizzeria': None,
        'Grocery Store': None,
        'Library': None
    }
    for i in final:
        add = []
        params = {
            "term": i,
            "location": f"{location}"
        }
        response = requests.get(url, headers=headers, params=params)
        try:
            businesses = response.json()["businesses"]
                
            rating = (sorted(businesses, key=lambda item: (item["rating"], (item["distance"]*-1))))
                #learn how to sort by distance!
            if len(rating) > 2:
                n1 = {
                    'name': rating[-1]["name"],
                    'location': (rating[-1]['location'])["address1"],
                    'rating': rating[-1]["rating"],
                    'phone': rating[-1]["phone"]
                }
                n2 = {
                    'name': rating[-2]["name"],
                    'location': (rating[-2]['location'])["address1"],
                    'rating': rating[-2]["rating"],
                    'phone': rating[-2]["phone"]
                }
                add.append(n1)
                add.append(n2)
        
            final[i] = add.copy()
        except:
            final[i] = 'nothing found'
        


    print(final)
    for location in final:
        if location != 'nothing found':
            return final
    return False


def isfloat(num):
    try:
        float(num)
        return True
    except ValueError:
        return False

def isint(num):
    try:
        int(num)
        return True
    except ValueError:
        return False

def GetText(link, look_at, SENTENCES_COUNT, country, university):
    url = link.split('%')[0]
    parser = HtmlParser.from_url(url, Tokenizer(LANGUAGE))
    # or for plain text files
    # parser = PlaintextParser.from_file("document.txt", Tokenizer(LANGUAGE))
    # parser = PlaintextParser.from_string("Check this out.", Tokenizer(LANGUAGE))
    stemmer = Stemmer(LANGUAGE)

    summarizer = Summarizer(stemmer)
    summarizer.stop_words = get_stop_words(LANGUAGE)
    final = []
    print(0)
    print(f"look at: {look_at} country: {country}")
    if look_at == 'needed+grades' and country == 'US':
        a = ScrapGoogle(university, '+university+average+gpa')
        a = a.split('All results')[-1]
        a = (a.split('. ')[0].strip() + '.').split(' ')
        for index, item in enumerate(a):
            if isfloat(item.strip(".")) or isint(item.strip(".")):
                a[index] += f' (or {round(float(item.strip("."))*5, 2)}/20 in france)'
                print('---------------------' + str(a))
                print('---------------------' + str(a[index]))
                a = ' '.join(a)      
                final.append(str(a))
                break
    
    for sentence in summarizer(parser.document, SENTENCES_COUNT):
        final.append(sentence)
    final = list(dict.fromkeys(final))
    return final, look_at, url




def ScrapGoogle(university, message):
    url = 'https://www.google.com/search?q=' + university + message
    print('looked up ' + url)
    headers = {"User-Agent": "Mozilla/5.0"}
    cookies = {"CONSENT": "YES+cb.20210720-07-p0.en+FX+410"}
    request_result = requests.get(url, headers=headers, cookies=cookies)
    soup = bs4.BeautifulSoup(request_result.text, "html.parser")
    texts = soup.findAll(text=True)
    visible_texts = filter(tag_visible, texts)  
    return u" ".join(t.strip() for t in visible_texts)
#idk why but only the second link with .edu works, the first one is weird
#try to remove the limit of lines you can print cuz i think they erase the first ones


def tag_visible(element):
    if not element.parent.name in ['div', 'span', 'b']:
        return False
    if isinstance(element, Comment):
        return False
    return True



LANGUAGE = "english"

def ReturnFirstURLs(university, item):
    URLs = []
    url = 'https://www.google.com/search?q=' + university + '+' + item
    headers = {"User-Agent": "Mozilla/5.0"}
    cookies = {"CONSENT": "YES+cb.20210720-07-p0.en+FX+410"}
    request_result = requests.get(url, headers=headers, cookies=cookies)
    soup = bs4.BeautifulSoup(request_result.text, "html.parser")
    heading_object = soup.find_all( 'a')
    for href in heading_object:
        URLs.append(href.get('href').split("/url?q=")[-1].split("&")[0])
    return URLs
    #idk why but only the second link with .edu works, the first one is weird
    #try to remove the limit of lines you can print cuz i think they erase the first ones

def filterLink(links):
    for item in links:
        if 'http' in item:
            if '.edu' in item and (not 'default/files/styles/' in item) and (not '.png' in item):
                #we should change the filter if it is about british unis i think they got .ac.uk
                return item


def filterLinkUK(links):
    for item in links:
        if 'http' in item:
            if '.ac.uk' in item and (not 'default/files/styles/' in item) and (not '.png' in item):
                #we should change the filter if it is about british unis i think they got .ac.uk
                return item


def DoForEach(university, SENTENCES_COUNT, list=['needed grades', 'application', 'cost'], country='US'):
    #major should be replaced by courses if it is a british uni
    #acceptance rate doesn't work, for ssome reason it is the link just after but we should learn how to take the number from the center of the web page
    returnn = []
    for item in list:
        if ' ' in item:
            item = '+'.join(item.split(' '))
        if country == 'UK':
            returnn.append(GetText(filterLinkUK(ReturnFirstURLs(university, item)), item, SENTENCES_COUNT, country, university))
        else:
            returnn.append(GetText(filterLink(ReturnFirstURLs(university, item)), item, SENTENCES_COUNT, country, university))
    return returnn







def check(text):
    months = ['january', 'jebruary', 'jarch', 'april', 'may', 'june', 'july', 'august', 'september', 'october', 'november', 'december']
    for sentence in text:
        for month in months:
            if month in sentence.lower():
                return sentence
    return 'no date'


def adFilter(text):
    if 'Ad' in text:
        return text.split('.')[-1]
    return text



def ScrapGoogle2(university, message):
    url = 'https://www.google.com/search?q=' + university + message
    print(url)
    headers = {"User-Agent": "Mozilla/5.0"}
    cookies = {"CONSENT": "YES+cb.20210720-07-p0.en+FX+410"}
    request_result = requests.get(url, headers=headers, cookies=cookies)
    soup = bs4.BeautifulSoup(request_result.text, "html.parser")
    texts = soup.findAll(text=True)
    visible_texts = filter(tag_visible2, texts)  
    return u" ".join(t.strip() for t in visible_texts)

def tag_visible2(element):
    if element.parent.name in ['style', 'script', 'head', 'title', 'meta', '[document]']:
        return False
    if isinstance(element, Comment):
        return False
    return True




























#-----------------------------------------------------------------------------------------------------------------------
 
# Flask constructor
app = Flask(__name__)

commments = [
    ["comment", "file"]
]
 
# A decorator used to tell the application
# which URL is associated function
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        comment = request.form.get("comment")
        fil = request.form.get("fil")
        if not comment:
            comment = '(No value)'
        if not fil:
            fil = '(No value)'
        with open("comments.csv", "a") as file:
            Writer = csv.writer(file)
            Writer.writerow([comment, fil])
    return render_template("index.html")


@app.route('/uksearch', methods=['GET', 'POST'])
def uksearch():
    name = request.args.get("name")
    listy = request.args.getlist("listy")
    SENTENCES_COUNT = request.args.get("lines")

    major = False
    sources = False
    if not SENTENCES_COUNT:
        SENTENCES_COUNT = 10
    if 'major' in listy:
        major = True
        listy.remove('major')
    else:
        major = False
    if 'sources' in listy:
        sources = True
        listy.remove('sources')
    else:
        sources = False
    if 'location' in listy:
        location = True
        listy.remove('location')
    else:
        location = False
    if 'acceptance rate' in listy:
        acceptance_rate = True
        listy.remove('acceptance rate')
    else:
        acceptance_rate = False

    
    

    uni = name.capitalize()

    all_text = list(DoForEach(uni, SENTENCES_COUNT, listy, 'UK'))
    for bob in all_text:
        for bobb in bob:
            print(bobb)
    if major:
        majors = filterLinkUK(ReturnFirstURLs(uni, 'major'))

    links = []
    headers = []
    text = []
    useStuffHead = []
    useStuff = []
    for item in list(all_text):
        headers.append(item[1].replace("+", " ").capitalize())

        text.append(item[0])
        links.append(item[-1])
    
    if acceptance_rate:
        a = ScrapGoogle(uni, '+university+acceptance+rate').split('All results')[-1]
        a = a.split('%')[0]
        a = a.strip() + '%'
        useStuffHead.append("Acceptance Rate")
        useStuff.append(a)
    
    if major:
        useStuffHead.append("Subjects")
        useStuff.append(majors)

    if sources:
        useStuffHead.append("Sources")
        useStuff.append(links)



    print(text[-1])
    print('\n\n\n' + str(SENTENCES_COUNT))
    print(listy)
    return render_template('greet.html', text=text, headers=headers, country='UK', uni=uni)







@app.route('/admin', methods=['GET', 'POST'])
def admin():
    if request.method == 'POST':
        items = []
        titles = []
        x = 0
        username = request.form.get("username")
        password = request.form.get("password")
        if (username == 'samuel' and password == 'verycomplexpassword') or (username == 'oskar' and password == 'password'):
            with open("comments.csv", "r") as file:
                reader = csv.reader(file)
                for row in reader:
                    print(row)
                    if x == 0:
                        titles = row.copy()
                    else:
                        print(1)
                        if row != []:
                            print(2)
                            items.append(row)
                    x += 1
            print(items)
            return render_template('admin.html', items=items, titles=titles)
    return render_template('admin.html')






@app.route('/ussearch')
def greet():
    name = request.args.get("name")
    listy = request.args.getlist("listy")
    SENTENCES_COUNT = request.args.get("lines")

    major = False
    sources = False
    if not SENTENCES_COUNT:
        SENTENCES_COUNT = 10
    if 'major' in listy:
        major = True
        listy.remove('major')
    else:
        major = False
    if 'sources' in listy:
        sources = True
        listy.remove('sources')
    else:
        sources = False
    if 'location' in listy:
        location = True
        listy.remove('location')
    else:
        location = False
    if 'acceptance rate' in listy:
        acceptance_rate = True
        listy.remove('acceptance rate')
    else:
        acceptance_rate = False

    uni = name.capitalize()

    all_text = list(DoForEach(uni, SENTENCES_COUNT, listy, 'US'))
    for bob in all_text:
        for bobb in bob:
            print(bobb)
    if major:
        majors = filterLink(ReturnFirstURLs(uni, 'major'))

    links = []
    headers = []
    text = []
    useStuffHead = []
    useStuff = []
    for item in list(all_text):
        headers.append(item[1].replace("+", " ").capitalize())

        text.append(item[0])
        links.append(item[-1])

    
    if major:
        useStuffHead.append("Majors")
        useStuff.append(majors)
    if acceptance_rate:
        a = ScrapGoogle(uni, '+university+acceptance+rate').split('All results')[-1]
        a = a.split('%')[0]
        a = a.strip() + '%'
        useStuffHead.append("Acceptance Rate")
        useStuff.append(a)


    try:
        try:
            deadline = ScrapGoogle2(uni, '+university+application+deadline+date').split('\n\n')[1]
            deadline = adFilter(check(deadline.replace('. ', ' .').replace('?', ' .').replace('›', ' .').replace('...', ' .').split(' .')).strip()).replace('\n', ' ') 
        except:
            deadline = ScrapGoogle2(uni, '+university+application+deadline+date').split('Verbatim')[1]
            deadline = adFilter(check(deadline.replace('. ', ' .').replace('?', ' .').replace('›', ' .').replace('...', ' .').split(' .')).strip()).replace('\n', ' ') 


        if deadline[-1] != '.':
            deadline += '.'

        useStuffHead.append('Deadline')
        useStuff.append(deadline)
    
    except:
        pass

    if location:
        print(ScrapGoogle(uni, '+university+location'))
        z = ScrapGoogle(uni, '+university+location').split('All results')[-1].split('- Wikipedia')[0].split('See results')[0].replace('›', '.').replace('|', '.').replace('/', '').replace('\\', '').split('.')
        print(z)
        for senty in z:   
            nearYou = PyYelp(senty.strip())
            print('--------------------' + str(nearYou))
            if nearYou:
                print(nearYou)
                useStuffHead.append("Location")
                useStuff.append(senty)
                useStuffHead.append("Near You")
                useStuff.append(nearYou)
                break



    if sources:
        useStuffHead.append("Sources")
        useStuff.append(links)
    
    doc = MakeDoc(headers, text, useStuff, useStuffHead, uni, 'US')
    useStuffHead.append("Document")
    useStuff.append(doc)



    print(text[-1])
    print(len(text))
    print('\n\n\n' + str(SENTENCES_COUNT))
    print(useStuff)
    return render_template('greet.html', text=text, headers=headers, country='US', uni=uni, useStuff=useStuff, useStuffHead=useStuffHead)
